import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import mysql.connector
from datetime import datetime
import re
from collections import Counter
from konlpy.tag import Okt
from config import MYSQL_CONFIG

class FashionAnalyzer:
    def __init__(self):
        self.okt = Okt()
        
    def fetch_data(self, query):
        """데이터베이스에서 데이터 가져오기"""
        try:
            with mysql.connector.connect(**MYSQL_CONFIG) as conn:
                cursor = conn.cursor()
                cursor.execute(query)
                columns = [column[0] for column in cursor.description]
                data = cursor.fetchall()
                return pd.DataFrame(data, columns=columns)
        except Exception as e:
            print(f"데이터 조회 오류: {str(e)}")
            return pd.DataFrame()
            
    def analyze_trend(self, period, season=None):
        """기간별 트렌드 분석"""
        today = datetime.now().strftime("%Y-%m-%d")
        query = f"""
        SELECT title, content, upload_date 
        FROM (
            SELECT * FROM vogue_trends
            UNION ALL
            SELECT * FROM wkorea_trends
            UNION ALL
            SELECT * FROM marieclaire_trends
            UNION ALL
            SELECT * FROM wwdkorea_trends
            UNION ALL
            SELECT * FROM jentestore_trends
        ) AS combined
        WHERE upload_date >= DATE_SUB('{today}', INTERVAL {period} DAY)
        """
        
        df = self.fetch_data(query)
        if df.empty:
            print("분석할 데이터가 없습니다.")
            return None
            
        # 텍스트 전처리
        texts = df['content'].fillna('').str.cat(sep=' ')
        words = self.okt.nouns(texts)
        word_count = Counter(words)
        
        # 패션 관련 불용어 추가
        stopwords = {'것', '등', '점', '수', '말', '거', '더', '전', '후', '중', '뭐', '때', '내'}
        fashion_keywords = {'패션', '스타일', '트렌드', '룩', '의류', '옷', '착용'}
        word_count = {k:v for k,v in word_count.items() 
                     if len(k) > 1 and k not in stopwords and (v > 1 or k in fashion_keywords)}
        
        if not word_count:
            print("키워드 추출 결과가 없습니다.")
            return None
        
        # 워드클라우드 생성
        wc = WordCloud(font_path='/System/Library/Fonts/AppleGothic.ttf',
                      background_color='white',
                      width=800, height=400)
        wc_img = wc.generate_from_frequencies(word_count)
        plt.figure(figsize=(10, 5))
        plt.imshow(wc_img)
        plt.axis('off')
        plt.savefig('trend_wordcloud.png')
        plt.close()
        
        # 트렌드 분석 추가
        trend_insights = self._analyze_trends(df)
        seasonal_items = self._get_seasonal_items(df)
        brand_mentions = self._analyze_brands(df)
        
        return {
            'top_keywords': dict(sorted(word_count.items(), key=lambda x: x[1], reverse=True)[:10]),
            'article_count': len(df),
            'period': period,
            'data': df,
            'trend_insights': trend_insights,
            'seasonal_items': seasonal_items,
            'brand_mentions': brand_mentions
        }

    def _analyze_trends(self, df):
        """주요 트렌드 분석"""
        trends = []
        content = ' '.join(df['content'].fillna(''))
        
        # 트렌드 패턴 분석
        trend_patterns = [
            r'(?:트렌드|유행하는|인기있는|핫한).*?(?:아이템|스타일|룩|패션).*?(?:은|는|임|입니다).*?([^.!?]*)',
            r'(?:이번|올해|이번\s*시즌).*?(?:트렌드|유행|스타일).*?(?:은|는|임|입니다).*?([^.!?]*)'
        ]
        
        for pattern in trend_patterns:
            matches = re.finditer(pattern, content)
            for match in matches:
                if trend := match.group(1).strip():
                    trends.append(trend)
        
        return trends[:5]  # 상위 5개 트렙드만 반환

    def _get_seasonal_items(self, df):
        """시즌별 주요 아이템 분석"""
        seasons = {
            '봄': ['봄', '스프링'],
            '여름': ['여름', '서머'],
            '가을': ['가을', '오터픔'],
            '겨울': ['겨울', '윈터']
        }
        
        seasonal_items = {season: [] for season in seasons}
        
        for season, keywords in seasons.items():
            season_content = df[df['content'].str.contains('|'.join(keywords), na=False)]
            if not season_content.empty:
                # 아이템 추출
                items = re.findall(r'([가-힣a-zA-Z]+\s*(?:재킷|코트|팬츠|스커트|원피스|셔츠|니트|백|슈즈|부츠))', 
                                 ' '.join(season_content['content']))
                seasonal_items[season] = list(set(items))[:5]
                
        return seasonal_items

    def _analyze_brands(self, df):
        """브랜드 언급 분석"""
        content = ' '.join(df['content'].fillna(''))
        brand_patterns = [
            r'(?:<([^>]+)>)',  # 태그로 둘러싸인 브랜드명
            r'(?:브랜드|하우스)\s+([가-힣a-zA-Z]+)',  # '브랜드' 다음에 나오는 이름
            r'([A-Z][a-zA-Z]+)\s+(?:컬렉션|룩|스타일)'  # 대문자로 시작하는 브랜드명
        ]
        
        brands = []
        for pattern in brand_patterns:
            matches = re.finditer(pattern, content)
            for match in matches:
                if brand := match.group(1).strip():
                    brands.append(brand)
                    
        return Counter(brands).most_common(5)

    def generate_report(self, analysis_result, report_type="트렌드 분석"):
        doc = Document()
        
        # 제목
        doc.add_heading(f'패션 트렌드 분석 리포트', 0)
        doc.add_paragraph(f'생성일: {datetime.now().strftime("%Y-%m-%d")}')
        doc.add_paragraph(f'분석 기간: 최근 {analysis_result["period"]}일')
        
        # 주요 통계
        doc.add_heading('주요 통계', level=1)
        doc.add_paragraph(f'분석 기사 수: {analysis_result["article_count"]}건')
        
        # 주요 키워드
        doc.add_heading('주요 키워드', level=1)
        keywords_para = doc.add_paragraph()
        for word, count in analysis_result['top_keywords'].items():
            keywords_para.add_run(f'{word}({count}회) ').bold = True
            
        # 워드클라우드 추가
        doc.add_heading('트렌드 워드클라우드', level=1)
        doc.add_picture('trend_wordcloud.png', width=Inches(6))
        
        # 트렌드 인사이트 추가
        doc.add_heading('주요 트렌드 인사이트', level=1)
        for idx, trend in enumerate(analysis_result['trend_insights'], 1):
            doc.add_paragraph(f'{idx}. {trend}')
            
        # 시즌별 아이템 추가
        doc.add_heading('시즌별 주요 아이템', level=1)
        for season, items in analysis_result['seasonal_items'].items():
            if items:
                p = doc.add_paragraph(f'{season}: ')
                p.add_run(', '.join(items)).italic = True
                
        # 인기 브랜드 추가
        doc.add_heading('가장 많이 언급된 브랜드', level=1)
        for brand, count in analysis_result['brand_mentions']:
            doc.add_paragraph(f'• {brand}: {count}회 언급')
        
        # 주요 기사 목록
        doc.add_heading('주요 기사', level=1)
        for _, row in analysis_result['data'].head(5).iterrows():
            p = doc.add_paragraph()
            p.add_run(f'{row["title"]}\n').bold = True
            p.add_run(f'작성일: {row["upload_date"]}\n')
            p.add_run(row["content"][:200] + "...\n\n")
        
        # 저장
        filename = f'패션트렌드분석_{datetime.now().strftime("%Y%m%d")}.docx'
        doc.save(filename)
        return filename

def main():
    analyzer = FashionAnalyzer()
    result = analyzer.analyze_trend(period=30)  # 최근 30일 분석
    if result:
        report_file = analyzer.generate_report(result)
        print(f"리포트가 생성되었습니다: {report_file}")

if __name__ == "__main__":
    main()
